from datetime import date
from types import SimpleNamespace
from io import BytesIO
from unittest import mock

from django.contrib.auth import get_user_model
from django.test import TestCase
from django.urls import reverse
from django.utils import timezone
from docx import Document as DocxDocument

from cadastros.models import Cargo
from cadastros.models import Cidade
from cadastros.models import Combustivel
from cadastros.models import Estado
from cadastros.models import Servidor
from cadastros.models import Unidade
from cadastros.models import Viatura
from documentos.services.types import DocumentoFormato
from oficios.models import Oficio
from roteiros.models import Roteiro
from roteiros.models import RoteiroDestino
from termos.forms import TermoAutorizacaoForm
from termos.models import TermoAutorizacao
from termos.services import VarianteTermo
from termos.services import build_termo_cadastro_payload
from termos.services import gerar_termo_cadastro_lote


class TermoAutorizacaoCadastroTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="tester_termo_cadastro", password="123456")
        self.client.force_login(self.user)
        self.cargo = Cargo.objects.create(nome="Investigador")
        self.unidade = Unidade.objects.create(nome="Unidade", sigla="UN")
        self.estado = Estado.objects.create(nome="Parana", sigla="PR")
        self.cidade = Cidade.objects.create(nome="Curitiba", estado=self.estado, uf="PR")
        self.cidade_destino = Cidade.objects.create(nome="Londrina", estado=self.estado, uf="PR")
        self.cidade_destino_2 = Cidade.objects.create(nome="Maringa", estado=self.estado, uf="PR")
        self.estado_sc = Estado.objects.create(nome="Santa Catarina", sigla="SC")
        self.cidade_sc = Cidade.objects.create(nome="Joinville", estado=self.estado_sc, uf="SC")
        self.servidor_1 = Servidor.objects.create(
            nome="Servidor Um",
            cargo=self.cargo,
            cpf="11111111111",
            rg="1234567",
            unidade=self.unidade,
        )
        self.servidor_2 = Servidor.objects.create(
            nome="Servidor Dois",
            cargo=self.cargo,
            cpf="22222222222",
            rg="7654321",
            unidade=self.unidade,
        )
        combustivel = Combustivel.objects.create(nome="Gasolina")
        self.viatura = Viatura.objects.create(
            placa="ABC1234",
            modelo="Duster",
            combustivel=combustivel,
            tipo=Viatura.TIPO_CARACTERIZADA,
        )
        self.roteiro = Roteiro.objects.create(
            origem_estado=self.estado,
            origem_cidade=self.cidade,
            saida_dt=timezone.make_aware(timezone.datetime(2026, 6, 10, 8, 0)),
            retorno_chegada_dt=timezone.make_aware(timezone.datetime(2026, 6, 12, 18, 0)),
        )
        self.destino_roteiro = RoteiroDestino.objects.create(
            roteiro=self.roteiro,
            estado=self.estado,
            cidade=self.cidade_destino,
            ordem=1,
        )
        self.oficio = Oficio.objects.create(
            numero=20,
            ano=2026,
            protocolo="123456781",
            assunto="Viagem",
            motivo="Evento",
            roteiro=self.roteiro,
            solicitante=self.unidade,
            viatura=self.viatura,
        )
        self.oficio.servidores.add(self.servidor_1)
        self.oficio.servidores_termo_autorizacao.add(self.servidor_1)

    def test_form_exige_destino_e_data_quando_nao_ha_oficio(self):
        form = TermoAutorizacaoForm(data={})
        self.assertFalse(form.is_valid())
        self.assertIn("destino_cidade", form.errors)
        self.assertIn("data_evento_inicio", form.errors)

    @mock.patch("termos.forms.resolver_sede_ids_desde_configuracao")
    def test_form_carrega_sede_padrao_da_configuracao(self, m_resolver):
        m_resolver.return_value = (self.estado.pk, self.cidade.pk, "")

        form = TermoAutorizacaoForm()

        self.assertEqual(form.fields["destino_estado"].initial, self.estado.pk)
        self.assertIsNone(form.fields["destino_cidade"].initial)
        self.assertEqual(form.fields["destino_estado"].widget.attrs["data-picker-initial-value"], str(self.estado.pk))
        self.assertEqual(form.fields["destino_cidade"].widget.attrs["data-picker-initial-value"], "")
        self.assertEqual(list(form.fields["destino_cidade"].queryset), [self.cidade, self.cidade_destino, self.cidade_destino_2])

    @mock.patch("termos.forms.resolver_sede_ids_desde_configuracao")
    def test_form_carrega_apenas_uf_quando_cidade_da_sede_nao_resolve(self, m_resolver):
        m_resolver.return_value = (self.estado.pk, None, "")

        form = TermoAutorizacaoForm()

        self.assertEqual(form.fields["destino_estado"].initial, self.estado.pk)
        self.assertIsNone(form.fields["destino_cidade"].initial)
        self.assertEqual(form.fields["destino_estado"].widget.attrs["data-picker-initial-value"], str(self.estado.pk))
        self.assertEqual(form.fields["destino_cidade"].widget.attrs["data-picker-initial-value"], "")
        self.assertEqual(list(form.fields["destino_cidade"].queryset), [self.cidade, self.cidade_destino, self.cidade_destino_2])

    def test_form_filtra_cidades_pela_uf_selecionada(self):
        form = TermoAutorizacaoForm(data={"destino_estado": str(self.estado_sc.pk)})

        cidade_ids = list(form.fields["destino_cidade"].queryset.values_list("pk", flat=True))
        self.assertEqual(cidade_ids, [self.cidade_sc.pk])

    def test_form_aceita_termo_avulso_sem_servidor_e_sem_viatura(self):
        form = TermoAutorizacaoForm(
            data={
                "destino_estado": str(self.estado.pk),
                "destino_cidade": str(self.cidade_destino.pk),
                "data_evento_inicio": "2026-06-10",
                "data_evento_fim": "",
                "servidores": [],
                "viatura": "",
            },
        )
        self.assertTrue(form.is_valid(), form.errors)
        termo = form.save()
        self.assertEqual(termo.data_evento_fim, date(2026, 6, 10))
        self.assertEqual(termo.servidores.count(), 0)
        self.assertIsNone(termo.viatura)

    def test_form_aceita_oficio_como_fallback_de_destino_e_periodo(self):
        form = TermoAutorizacaoForm(data={"oficio": str(self.oficio.pk)})
        self.assertTrue(form.is_valid(), form.errors)

    def test_periodo_invalido_retorna_erro(self):
        form = TermoAutorizacaoForm(
            data={
                "destino_estado": str(self.estado.pk),
                "destino_cidade": str(self.cidade_destino.pk),
                "data_evento_inicio": "2026-06-12",
                "data_evento_fim": "2026-06-10",
            },
        )
        self.assertFalse(form.is_valid())
        self.assertIn("data_evento_fim", form.errors)

    def test_vinculo_com_oficio_permanece_vivo_para_destino(self):
        termo = TermoAutorizacao.objects.create(oficio=self.oficio)
        self.assertIn("LONDRINA", termo.destino_display)

        self.destino_roteiro.cidade = self.cidade_destino_2
        self.destino_roteiro.save()
        termo = TermoAutorizacao.objects.get(pk=termo.pk)

        self.assertIn("MARINGA", termo.destino_display)

    def test_campos_do_termo_tem_precedencia_sobre_oficio(self):
        termo = TermoAutorizacao.objects.create(
            oficio=self.oficio,
            destino_estado=self.estado,
            destino_cidade=self.cidade_destino_2,
            data_evento_inicio=date(2026, 7, 1),
            data_evento_fim=date(2026, 7, 1),
        )

        self.assertIn("MARINGA", termo.destino_display)
        self.assertEqual(termo.periodo_display, "01/07/2026")

    def test_servidores_do_termo_substituem_fallback_do_oficio(self):
        termo = TermoAutorizacao.objects.create(oficio=self.oficio)
        self.assertEqual(list(termo.servidores_efetivos()), [self.servidor_1])

        termo.servidores.add(self.servidor_2)

        self.assertEqual(list(termo.servidores_efetivos()), [self.servidor_2])

    def test_termo_generico_do_evento_usa_servidores_dos_oficios(self):
        from eventos.models import Evento

        evento = Evento.objects.create(titulo="Acao", data_inicio=date(2026, 6, 10))
        self.oficio.evento = evento
        self.oficio.save(update_fields=["evento"])
        self.oficio.servidores_termo_autorizacao.add(self.servidor_2)

        termo_generico = TermoAutorizacao.objects.create(
            evento=evento,
            destino_estado=self.estado,
            destino_cidade=self.cidade_destino,
            data_evento_inicio=date(2026, 6, 10),
            data_evento_fim=date(2026, 6, 10),
        )

        self.assertEqual(
            list(termo_generico.servidores_efetivos()),
            [self.servidor_2, self.servidor_1],
        )

    def test_payload_sem_servidor_e_semipreenchido(self):
        termo = TermoAutorizacao.objects.create(
            destino_estado=self.estado,
            destino_cidade=self.cidade_destino,
            data_evento_inicio=date(2026, 6, 10),
            data_evento_fim=date(2026, 6, 10),
        )

        payload = build_termo_cadastro_payload(termo, None)

        self.assertEqual(payload["termo"]["variante"], VarianteTermo.SEMIPREENCHIDO)
        self.assertEqual(payload["termo"]["participante"]["nome"], "")
        self.assertIn("LONDRINA", payload["termo"]["viagem"]["destino_principal"])

    @mock.patch("termos.services.gerar_termo_cadastro_um")
    def test_lote_gera_um_documento_por_servidor(self, m_gerar):
        termo = TermoAutorizacao.objects.create(
            destino_estado=self.estado,
            destino_cidade=self.cidade_destino,
            data_evento_inicio=date(2026, 6, 10),
            data_evento_fim=date(2026, 6, 10),
        )
        termo.servidores.set([self.servidor_1, self.servidor_2])
        m_gerar.return_value = SimpleNamespace(conteudo=b"x", nome_arquivo="x.docx")

        docs = gerar_termo_cadastro_lote(termo, DocumentoFormato.DOCX)

        self.assertEqual(len(docs), 2)
        self.assertEqual(m_gerar.call_count, 2)

    def test_post_cria_termo_avulso(self):
        response = self.client.post(
            reverse("termos:novo"),
            data={
                "destino_estado": str(self.estado.pk),
                "destino_cidade": str(self.cidade_destino.pk),
                "data_evento_inicio": "2026-06-10",
                "data_evento_fim": "2026-06-10",
                "servidores": [str(self.servidor_1.pk), str(self.servidor_2.pk)],
                "viatura": "",
                "action": "save",
            },
        )

        self.assertEqual(response.status_code, 302)
        termo = TermoAutorizacao.objects.get()
        self.assertEqual(termo.servidores.count(), 2)

    def test_get_index_e_novo_renderizam(self):
        response_index = self.client.get(reverse("termos:index"))
        self.assertEqual(response_index.status_code, 200)
        self.assertContains(response_index, "Novo termo")

        response_novo = self.client.get(reverse("termos:novo"))
        self.assertEqual(response_novo.status_code, 200)
        self.assertContains(response_novo, "Cadastro de termo")
        self.assertNotContains(response_novo, "page-stepper page-stepper--horizontal")
        self.assertContains(response_novo, "id=\"id_oficio_busca\"")
        self.assertContains(response_novo, "id=\"termo-oficio-lista\"")
        self.assertContains(response_novo, "id=\"termo-evento-date-picker\"")
        self.assertContains(response_novo, "id=\"termo-btn-adicionar-destino\"")

        termo = TermoAutorizacao.objects.create(
            destino_estado=self.estado,
            destino_cidade=self.cidade_destino,
            data_evento_inicio=date(2026, 6, 10),
            data_evento_fim=date(2026, 6, 10),
        )
        response_edit = self.client.get(reverse("termos:editar", args=[termo.pk]))
        self.assertEqual(response_edit.status_code, 200)
        self.assertNotContains(response_edit, "page-stepper page-stepper--horizontal")
        self.assertContains(response_edit, "id=\"id_oficio_busca\"")

    @staticmethod
    def _docx_bytes(texto):
        doc = DocxDocument()
        doc.add_paragraph(texto)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _fake_gerar_docx(self, content_type):
        def _fake(termo, servidor, formato):
            if servidor is None:
                texto = "GENERICO"
            elif servidor.pk == self.servidor_1.pk:
                texto = "SERVIDOR A"
            else:
                texto = "SERVIDOR B"
            return SimpleNamespace(
                conteudo=self._docx_bytes(texto),
                nome_arquivo="x.docx",
                content_type=content_type,
                hash_sha256=str(getattr(servidor, "pk", 0)),
            )

        return _fake

    @mock.patch("termos.async_documents.gerar_termo_cadastro_um")
    def test_download_docx_multiplo_retorna_arquivo_unico(self, m_um):
        termo = TermoAutorizacao.objects.create(
            destino_estado=self.estado,
            destino_cidade=self.cidade_destino,
            data_evento_inicio=date(2026, 6, 10),
            data_evento_fim=date(2026, 6, 10),
        )
        termo.servidores.set([self.servidor_1, self.servidor_2])
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        m_um.side_effect = self._fake_gerar_docx(content_type)

        response = self.client.get(reverse("termos:baixar_termo_cadastro_docx", args=[termo.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], content_type)
        self.assertEqual(response["Content-Disposition"], f'attachment; filename="termo_{termo.pk}.docx"')

        merged = DocxDocument(BytesIO(b"".join(response.streaming_content)))
        textos = [p.text for p in merged.paragraphs]
        self.assertIn("GENERICO", textos)
        self.assertIn("SERVIDOR A", textos)
        self.assertIn("SERVIDOR B", textos)

    @mock.patch("termos.async_documents.gerar_termo_cadastro_um")
    def test_download_sem_servidor_nao_duplica_termo_generico(self, m_um):
        """Termo sem servidores efetivos gera UM genérico apenas (sem página repetida)."""
        termo = TermoAutorizacao.objects.create(
            destino_estado=self.estado,
            destino_cidade=self.cidade_destino,
            data_evento_inicio=date(2026, 6, 10),
            data_evento_fim=date(2026, 6, 10),
        )
        m_um.side_effect = self._fake_gerar_docx(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        self.client.get(reverse("termos:baixar_termo_cadastro_docx", args=[termo.pk]))

        self.assertEqual(m_um.call_count, 1)
        (_, servidor_arg, _), _ = m_um.call_args
        self.assertIsNone(servidor_arg)
