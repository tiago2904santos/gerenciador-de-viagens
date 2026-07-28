"""Golden files dos templates documentais (N-04).

Antes destes testes, **nenhum teste abria o arquivo produzido**: os testes do
facade mockam ``_render_docx`` e afirmam sobre bytes fabricados. Com 5 motores
de PDF intercambiáveis e 11 templates versionados como binário, uma alteração
acidental num ``.docx`` passava por toda a suíte sem ser notada.

O que estes testes garantem
---------------------------
Para cada template de ``documentos/resources/``: o arquivo abre, renderiza, e o
**texto resultante** é exatamente o registrado no golden correspondente. Isso
trava três coisas que hoje ninguém vigia:

* o texto fixo do documento (cabeçalho, rodapé, cláusulas, rótulos de coluna);
* o conjunto de variáveis que o template consome e **onde** cada uma cai;
* a integridade do binário — um ``.docx`` corrompido ou trocado não renderiza.

O que estes testes NÃO garantem
-------------------------------
Que o *valor* enviado a cada variável esteja certo. O contexto aqui é sintético
(cada variável recebe ``[[nome]]``), de propósito: isolar o template do
construtor de contexto. A correção dos valores é responsabilidade dos testes de
cada módulo — é lá que ``oficios/docxtpl_context.py`` e companhia são exercidos.

Atualizando um golden
---------------------
Alterou um template de propósito? Rode::

    python manage.py test documentos.tests.test_golden_templates --settings=config.settings.test

e, com a falha em mãos, regrave o golden com ``ATUALIZAR_GOLDEN=1``::

    ATUALIZAR_GOLDEN=1 python manage.py test documentos.tests.test_golden_templates ...

Reveja o diff do golden no PR: ele é a evidência do que mudou no documento.
"""

from __future__ import annotations

import io
import os
from pathlib import Path

from django.test import SimpleTestCase

RESOURCES = Path(__file__).resolve().parents[1] / "resources"
GOLDEN = Path(__file__).resolve().parent / "golden"

# Marcador propositalmente sem "<" nem ">": docxtpl injeta o valor cru no XML do
# documento, e caracteres de marcação quebram o arquivo antes de renderizar.
MARCADOR = "[[{}]]"


def _contexto_sintetico(variaveis) -> dict[str, str]:
    return {nome: MARCADOR.format(nome) for nome in sorted(variaveis)}


def _texto_do_docx(conteudo: bytes) -> str:
    """Texto do documento na ordem em que aparece, parágrafos e tabelas."""
    from docx import Document

    doc = Document(io.BytesIO(conteudo))
    linhas: list[str] = []
    for paragrafo in doc.paragraphs:
        texto = paragrafo.text.strip()
        if texto:
            linhas.append(texto)
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells]
            if any(celulas):
                linhas.append(" | ".join(celulas))
    return "\n".join(linhas) + "\n"


def _renderizar(caminho: Path) -> str:
    from docxtpl import DocxTemplate

    template = DocxTemplate(str(caminho))
    template.render(_contexto_sintetico(template.get_undeclared_template_variables()))
    buffer = io.BytesIO()
    template.save(buffer)
    return _texto_do_docx(buffer.getvalue())


class GoldenTemplatesDocxTests(SimpleTestCase):
    """Um teste por template, para a falha apontar o arquivo culpado."""

    maxDiff = None

    @classmethod
    def templates(cls) -> list[Path]:
        return sorted(RESOURCES.glob("*.docx"))

    def _comparar(self, template: Path) -> None:
        atual = _renderizar(template)
        golden = GOLDEN / f"{template.stem}.txt"

        if os.environ.get("ATUALIZAR_GOLDEN") == "1":
            golden.parent.mkdir(parents=True, exist_ok=True)
            golden.write_text(atual, encoding="utf-8")
            self.skipTest(f"golden regravado: {golden.name}")

        self.assertTrue(
            golden.exists(),
            f"Falta o golden de {template.name}. Gere com ATUALIZAR_GOLDEN=1 e revise o diff.",
        )
        self.assertEqual(
            atual,
            golden.read_text(encoding="utf-8"),
            f"O texto renderizado de {template.name} mudou. Se foi de propósito, "
            f"regrave o golden com ATUALIZAR_GOLDEN=1 e revise o diff no PR.",
        )

    def test_todos_os_templates_tem_golden(self):
        """Template novo sem golden reprova — senão ele entra sem vigilância."""
        sem_golden = [
            t.name for t in self.templates() if not (GOLDEN / f"{t.stem}.txt").exists()
        ]
        self.assertEqual(sem_golden, [], f"Templates sem golden: {sem_golden}")


class GoldenDiarioBordoXlsxTests(SimpleTestCase):
    """O diário de bordo é o 11º documento, e o único em planilha.

    Vale mais que os demais: ``fill_diario_bordo_xlsx`` não substitui texto, ele
    **escreve em coordenadas fixas** (``TRECHO_COLUNAS``, ``TRECHO_ROW``). Se
    alguém inserir uma coluna no modelo, nada quebra em tempo de execução — os
    dados passam a cair na célula errada, em silêncio. O golden pega isso.
    """

    maxDiff = None
    TEMPLATE = RESOURCES / "diario_bordo.xlsx"
    GOLDEN_PATH = GOLDEN / "diario_bordo.txt"

    def _renderizar(self) -> str:
        import openpyxl

        from documentos.services.adapters.xlsx_render import (
            TRECHO_COLUNAS,
            fill_diario_bordo_xlsx,
        )

        header = {
            "divisao": "[[divisao]]",
            "placa": "[[placa]]",
            "motorista": "[[motorista]]",
        }
        trechos = [
            {campo: f"[[t{i}.{campo}]]" for campo in TRECHO_COLUNAS} for i in (1, 2)
        ]
        conteudo = fill_diario_bordo_xlsx(
            template_path=self.TEMPLATE, header=header, trechos=trechos
        )

        ws = openpyxl.load_workbook(io.BytesIO(conteudo)).active
        linhas = []
        for linha in ws.iter_rows(values_only=True):
            celulas = ["" if c is None else str(c).strip() for c in linha]
            if any(celulas):
                linhas.append(" | ".join(celulas).rstrip(" |"))
        return "\n".join(linhas) + "\n"

    def test_golden_diario_bordo(self):
        atual = self._renderizar()
        if os.environ.get("ATUALIZAR_GOLDEN") == "1":
            self.GOLDEN_PATH.parent.mkdir(parents=True, exist_ok=True)
            self.GOLDEN_PATH.write_text(atual, encoding="utf-8")
            self.skipTest("golden regravado: diario_bordo.txt")

        self.assertTrue(self.GOLDEN_PATH.exists(), "Falta o golden do diário de bordo.")
        self.assertEqual(
            atual,
            self.GOLDEN_PATH.read_text(encoding="utf-8"),
            "A planilha do diário de bordo mudou. Confira se alguma coluna foi "
            "inserida: as coordenadas de TRECHO_COLUNAS são fixas.",
        )

    def test_dois_trechos_produzem_duas_linhas(self):
        """Cada trecho vira uma linha própria — o modelo tem só a linha-gabarito."""
        texto = self._renderizar()
        self.assertIn("[[t1.km_inicial]]", texto)
        self.assertIn("[[t2.km_inicial]]", texto)


def _adicionar_testes() -> None:
    """Cria um método de teste por template, nomeado pelo arquivo."""

    for template in GoldenTemplatesDocxTests.templates():
        nome = f"test_golden_{template.stem.replace('-', '_')}"

        def teste(self, _template=template):
            self._comparar(_template)

        teste.__name__ = nome
        teste.__doc__ = f"O texto renderizado de {template.name} é o do golden."
        setattr(GoldenTemplatesDocxTests, nome, teste)


_adicionar_testes()
